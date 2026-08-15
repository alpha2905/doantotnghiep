#!/usr/bin/env python3
"""
Script tự động thêm diagrams vào file Word
Yêu cầu: pip install python-docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_image_with_caption(doc, image_path, caption):
    """Thêm hình ảnh vào Word document với caption"""
    
    if not os.path.exists(image_path):
        print(f"⚠️  Không tìm thấy: {image_path}")
        # Thêm placeholder
        p = doc.add_paragraph(f"[{caption} - File không tìm thấy]")
        p.runs[0].italic = True
        return False
    
    try:
        # Thêm hình ảnh
        doc.add_picture(image_path, width=Inches(5.5))
        
        # Thêm caption
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Khoảng trống
        
        return True
    except Exception as e:
        print(f"❌ Lỗi khi chèn {image_path}: {e}")
        return False

def update_word_with_diagrams():
    """Cập nhật file Word với diagrams"""
    
    # File Word input/output
    word_file = 'Chuong3_PhanTichThietKe.docx'
    output_file = 'Chuong3_PhanTichThietKe_VoiHinh.docx'
    
    if not os.path.exists(word_file):
        print(f"❌ File không tìm thấy: {word_file}")
        return False
    
    print(f"📖 Mở file: {word_file}")
    doc = Document(word_file)
    
    # Folder chứa diagrams
    diagrams_dir = Path('diagrams/images')
    
    # Danh sách diagrams cần thêm
    diagrams = [
        {
            'file': diagrams_dir / 'context.puml.png',
            'caption': 'Hình 3-1: Sơ đồ ngữ cảnh (Context Diagram)',
            'position': '3.1.2',  # Section để chèn
        },
        {
            'file': diagrams_dir / 'usecase.puml.png',
            'caption': 'Hình 3-2: Sơ đồ Use Case chính',
            'position': '3.1.3',
        },
        {
            'file': diagrams_dir / 'architecture.puml.png',
            'caption': 'Hình 3-3: Sơ đồ kiến trúc hệ thống phân tầng',
            'position': '3.3.1',
        },
        {
            'file': diagrams_dir / 'entity_relationship.puml.png',
            'caption': 'Hình 3-4: Sơ đồ ER MongoDB',
            'position': '3.3.2',
        },
        {
            'file': diagrams_dir / 'sequence_search.puml.png',
            'caption': 'Hình 3-6: Sơ đồ tuần tự Use Case Tìm Kiếm Sản Phẩm',
            'position': '3.4.1',
        },
        {
            'file': diagrams_dir / 'sequence_crawl.puml.png',
            'caption': 'Hình 3-7: Sơ đồ tuần tự Use Case Crawl Dữ Liệu',
            'position': '3.4.2',
        },
        {
            'file': diagrams_dir / 'sequence_sentiment.puml.png',
            'caption': 'Hình 3-8: Sơ đồ tuần tự Use Case Phân Tích Cảm Xúc',
            'position': '3.4.3',
        },
        {
            'file': diagrams_dir / 'sequence_forecast.puml.png',
            'caption': 'Hình 3-9: Sơ đồ tuần tự Use Case Dự Báo Giá',
            'position': '3.4.4',
        },
    ]
    
    print(f"\n📁 Folder diagrams: {diagrams_dir}")
    print(f"{'='*60}")
    
    # Cách 1: Thêm ở cuối document (đơn giản nhất)
    print("\n💡 Phương pháp: Thêm diagrams ở cuối document")
    print("   (Bạn có thể tự di chuyển chúng đến vị trí cần thiết trong Word)")
    
    doc.add_page_break()
    
    # Title cho section diagrams
    heading = doc.add_heading('PHỤ LỤC: CÁC SƠ ĐỒ CHI TIẾT', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Thêm các diagrams
    success_count = 0
    for diagram in diagrams:
        image_path = diagram['file']
        caption = diagram['caption']
        
        print(f"\n⏳ Xử lý: {caption}")
        
        if insert_image_with_caption(doc, str(image_path), caption):
            print(f"   ✅ Thêm thành công")
            success_count += 1
        else:
            print(f"   ⚠️  Bỏ qua (file không tìm thấy)")
    
    # Lưu file
    print(f"\n{'='*60}")
    print(f"💾 Lưu file: {output_file}")
    
    try:
        doc.save(output_file)
        print(f"✅ Lưu thành công!")
        print(f"\n📊 Kết quả:")
        print(f"   ✅ Diagrams thêm thành công: {success_count}/{len(diagrams)}")
        print(f"   📄 File output: {output_file}")
        
        return True
    except Exception as e:
        print(f"❌ Lỗi khi lưu: {e}")
        return False

def main():
    """Main function"""
    print("\n" + "="*60)
    print("🔧 TOOL: Tự động thêm Diagrams vào file Word")
    print("="*60)
    
    print("\n📋 Bước 1: Kiểm tra các file cần thiết...")
    
    # Kiểm tra files
    if not os.path.exists('Chuong3_PhanTichThietKe.docx'):
        print("❌ Chưa tìm thấy file: Chuong3_PhanTichThietKe.docx")
        print("💡 Chạy lệnh trước: python generate_chapter3.py")
        return
    
    if not os.path.exists('diagrams/images'):
        print("⚠️  Folder diagrams/images chưa tồn tại")
        print("💡 Bước tiếp theo:")
        print("   1. Chạy: python convert_diagrams.py")
        print("   2. Hoặc tạo folder: mkdir -p diagrams/images")
        print("   3. Thêm PNG files vào folder này")
        print("\n💻 Hãy chuẩn bị diagrams trước, rồi chạy lại script này")
        return
    
    print("✅ Các files cần thiết đã tìm thấy")
    
    # Cập nhật Word
    print("\n📝 Bước 2: Cập nhật file Word...")
    if update_word_with_diagrams():
        print("\n" + "="*60)
        print("🎉 Hoàn tất! File Word đã được cập nhật với diagrams")
        print("="*60)
    else:
        print("\n❌ Có lỗi xảy ra khi cập nhật file Word")

if __name__ == '__main__':
    main()
